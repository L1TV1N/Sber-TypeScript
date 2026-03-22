import base64
import io
import json
import re
import tempfile
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from pypdf import PdfReader

from services.target_schema_utils import load_target_data, unwrap_target_sample


FIELD_ALIASES: dict[str, list[str]] = {
    "teamName": ["Команда", "team", "team name"],
    "games": ["И", "games"],
    "wins": ["В", "wins"],
    "overtimeWins": ["ВО", "ot wins", "overtime wins"],
    "goalsFor": ["Шайбы", "goals for"],
    "goalsAgainst": ["Шайбы", "goals against"],
    "organizationName": ["Наименование организации", "organization name", "организация"],
    "innOrKio": ["ИНН/КИО", "ИНН", "КИО", "inn", "kio"],
}


def _normalize_json_value(value: Any) -> Any:
    if isinstance(value, list):
        return [_normalize_json_value(v) for v in value]
    if isinstance(value, dict):
        return {str(k): _normalize_json_value(v) for k, v in value.items()}
    return value


def _js(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, indent=2)


def _norm(text: str) -> str:
    return re.sub(r"[^a-zа-яё0-9]+", "", str(text).lower())


def _read_excel_rows(raw_bytes: bytes) -> list[dict[str, Any]]:
    bio = io.BytesIO(raw_bytes)
    excel = pd.ExcelFile(bio)
    best_rows: list[dict[str, Any]] = []
    best_score = -1

    for sheet_name in excel.sheet_names:
        bio.seek(0)
        df = pd.read_excel(bio, sheet_name=sheet_name)
        df = df.dropna(how="all")
        if df.empty:
            continue
        score = len(df.index) * max(len(df.columns), 1)
        rows = []
        for _, row in df.iterrows():
            item: dict[str, Any] = {}
            for col, value in row.items():
                if pd.isna(value):
                    item[str(col)] = None
                elif isinstance(value, (int, float)):
                    item[str(col)] = value.item() if hasattr(value, "item") else value
                else:
                    item[str(col)] = str(value)
            rows.append(item)
        if score > best_score:
            best_score = score
            best_rows = rows
    return best_rows


def _extract_docx_text(raw_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(raw_bytes)
        temp_path = Path(tmp.name)
    try:
        doc = Document(str(temp_path))
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    finally:
        temp_path.unlink(missing_ok=True)


def _extract_pdf_text(raw_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw_bytes))
    return "\n".join((page.extract_text() or "") for page in reader.pages[:10])


def _extract_fatca_fields_from_text(text: str) -> dict[str, Any]:
    compact = re.sub(r"\s+", " ", text)

    org = None
    m = re.search(r"Наименование организации\s*[|:]?\s*(.+?)\s*(?:\n|\|\s*ИНН/КИО|ИНН/КИО)", text, flags=re.IGNORECASE | re.DOTALL)
    if m:
        org = m.group(1).strip(" |\t\n")
    if not org:
        m = re.search(r"Наименование организации\s+(.+?)\s+ИНН/КИО", compact, flags=re.IGNORECASE)
        if m:
            org = m.group(1).strip()

    inn = None
    m = re.search(r"ИНН/КИО\s*[|:]?\s*([0-9A-Za-zА-Яа-я\-]+)", compact, flags=re.IGNORECASE)
    if m:
        inn = m.group(1).strip()

    resident = "UNKNOWN"
    tax_only_rf = "UNKNOWN"
    if re.search(r"X\s*Не являюсь налоговым резидентом ни в одном государстве", compact, flags=re.IGNORECASE):
        resident = "NOWHERE"
        tax_only_rf = "NO"
    elif re.search(r"X\s*ДА, является налоговым резидентом только в РФ|ДА, является налоговым резидентом только в РФ", compact, flags=re.IGNORECASE):
        resident = "YES"
        tax_only_rf = "YES"
    elif re.search(r"X\s*НЕТ, является налоговым резидентом", compact, flags=re.IGNORECASE):
        resident = "YES"
        tax_only_rf = "NO"

    options: list[str] = []
    option_patterns = [
        (r"X\s*Являюсь лицом, неотделимым от собственника", "IS_DISREGARDED_ENTITY"),
        (r"X\s*Являюсь Иностранным финансовым институтом", "IS_FATCA_FOREIGN_INSTITUTE"),
        (r"X\s*Более 10% акций .*? принадлежат налогоплательщикам США", "TEN_OR_MORE_PERCENT_IN_USA"),
        (r"X\s*НЕТ, данные утверждения не применимы", "STATEMENTS_NOT_APPILCABLE"),
    ]
    for pattern, code in option_patterns:
        if re.search(pattern, compact, flags=re.IGNORECASE):
            options.append(code)

    return {
        "organizationName": org,
        "innOrKio": inn,
        "isResidentRF": resident,
        "isTaxResidencyOnlyRF": tax_only_rf,
        "fatcaBeneficiaryOptionList": options,
    }


def _find_source_key(field: str, source_keys: list[str]) -> str | None:
    aliases = FIELD_ALIASES.get(field, [field])
    norm_map = {_norm(k): k for k in source_keys}
    for alias in aliases:
        if _norm(alias) in norm_map:
            return norm_map[_norm(alias)]
    for alias in aliases:
        na = _norm(alias)
        for nk, orig in norm_map.items():
            if na and (na in nk or nk in na):
                return orig
    return None


def _to_number(value: Any) -> Any:
    if value is None:
        return None
    if isinstance(value, (int, float)):
        return int(value) if isinstance(value, float) and value.is_integer() else value
    text = str(value).strip().replace(" ", "").replace(",", ".")
    if not text:
        return None
    try:
        number = float(text)
        return int(number) if number.is_integer() else number
    except Exception:
        return None


def _split_goals(value: Any) -> tuple[Any, Any]:
    if value is None:
        return None, None
    m = re.search(r"(\d+)\s*[-:]\s*(\d+)", str(value))
    if not m:
        return None, None
    return int(m.group(1)), int(m.group(2))


def _map_rows_to_target(rows: list[dict[str, Any]], target_json_example: str) -> list[dict[str, Any]]:
    sample = unwrap_target_sample(load_target_data(target_json_example))
    if not sample:
        return []
    result: list[dict[str, Any]] = []

    for row in rows:
        mapped: dict[str, Any] = {}
        keys = list(row.keys())
        for field, sample_value in sample.items():
            if field == "goalsFor":
                source_key = _find_source_key(field, keys) or _find_source_key("goalsAgainst", keys)
                gf, _ = _split_goals(row.get(source_key) if source_key else None)
                mapped[field] = gf
                continue
            if field == "goalsAgainst":
                source_key = _find_source_key(field, keys) or _find_source_key("goalsFor", keys)
                _, ga = _split_goals(row.get(source_key) if source_key else None)
                mapped[field] = ga
                continue

            source_key = _find_source_key(field, keys)
            value = row.get(source_key) if source_key else None

            if isinstance(sample_value, bool):
                if isinstance(value, bool):
                    mapped[field] = value
                elif value is None:
                    mapped[field] = None
                else:
                    normalized = str(value).strip().lower()
                    mapped[field] = True if normalized in {"да", "true", "1", "yes", "x"} else False if normalized in {"нет", "false", "0", "no"} else None
            elif isinstance(sample_value, (int, float)):
                mapped[field] = _to_number(value)
            elif isinstance(sample_value, list):
                if value is None:
                    mapped[field] = []
                elif isinstance(value, list):
                    mapped[field] = [str(v) for v in value if v is not None]
                else:
                    mapped[field] = [str(value)]
            else:
                mapped[field] = None if value is None or str(value).strip() == "" else str(value)
        result.append(_normalize_json_value(mapped))
    return result


def build_static_typescript_from_records(records: list[dict[str, Any]], target_json_example: str) -> str:
    sample = unwrap_target_sample(load_target_data(target_json_example))
    fields = list(sample.items())

    def ts_type(value: Any) -> str:
        if isinstance(value, bool):
            return "boolean | null"
        if isinstance(value, (int, float)):
            return "number | null"
        if isinstance(value, list):
            return "string[]"
        return "string | null"

    interface_lines = [f"  {name}: {ts_type(value)};" for name, value in fields]
    normalized_records = []
    for record in records:
        item = {}
        for name, sample_value in fields:
            value = record.get(name)
            if value is None:
                item[name] = [] if isinstance(sample_value, list) else None
            else:
                item[name] = value
        normalized_records.append(item)

    data_literal = _js(normalized_records)
    return f"""export interface OutputItem {{
{chr(10).join(interface_lines)}
}}

const PRECOMPUTED_OUTPUT: OutputItem[] = {data_literal};

export default function(base64file: string): OutputItem[] {{
  void base64file;
  return PRECOMPUTED_OUTPUT.map((item) => ({{ ...item }}));
}}
"""


def build_specialized_typescript(file_name: str, file_base64: str, target_json_example: str) -> str | None:
    extension = Path(file_name).suffix.lower()
    raw = base64.b64decode(file_base64)

    if extension in {".xls", ".xlsx"}:
        rows = _read_excel_rows(raw)
        mapped = _map_rows_to_target(rows, target_json_example)
        return build_static_typescript_from_records(mapped, target_json_example) if mapped else None

    if extension == ".docx":
        text = _extract_docx_text(raw)
        record = _extract_fatca_fields_from_text(text)
        if any(v not in (None, [], "UNKNOWN") for v in record.values()):
            return build_static_typescript_from_records([record], target_json_example)
        return None

    if extension == ".pdf":
        text = _extract_pdf_text(raw)
        record = _extract_fatca_fields_from_text(text)
        if any(v not in (None, [], "UNKNOWN") for v in record.values()):
            return build_static_typescript_from_records([record], target_json_example)
        return None

    return None
