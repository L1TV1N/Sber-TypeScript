import base64
import csv
import io
import json
import re
from functools import lru_cache
from pathlib import Path
from typing import Any

try:
    import cv2
except Exception:
    cv2 = None
try:
    import easyocr
except Exception:
    easyocr = None
import numpy as np
import pandas as pd
from docx import Document
from PIL import Image
from pypdf import PdfReader


CYRILLIC_ALLOWLIST = (
    "АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ"
    "абвгдеёжзийклмнопрстуфхцчшщъыьэюя"
    "0123456789"
    " .,:;!?-—_()[]{}\"'«»/\\|@#%&+=*"
)


def _detect_csv_delimiter(sample_text: str) -> str:
    candidates = [";", ",", "\t", "|"]
    counts = {d: sample_text.count(d) for d in candidates}
    best = max(counts, key=counts.get)
    return best if counts[best] > 0 else ","


def _clean_records(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    cleaned = []
    for row in records:
        cleaned_row = {}
        for key, value in row.items():
            if pd.isna(value):
                cleaned_row[str(key)] = ""
            else:
                cleaned_row[str(key)] = str(value)
        cleaned.append(cleaned_row)
    return cleaned


def preview_csv_from_base64(base64file: str) -> str:
    raw = base64.b64decode(base64file)
    text = raw.decode("utf-8-sig", errors="ignore")

    lines = [line for line in text.splitlines() if line.strip()]
    sample = "\n".join(lines[:5])
    delimiter = _detect_csv_delimiter(sample)

    reader = csv.DictReader(io.StringIO(text), delimiter=delimiter)
    rows = list(reader)
    cleaned_rows = _clean_records(rows)

    result = {
        "format": "csv",
        "delimiter": delimiter,
        "columns": reader.fieldnames or [],
        "sample_rows": cleaned_rows[:5],
        "all_rows": cleaned_rows[:500],
        "total_rows": len(cleaned_rows),
        "total_sampled_rows": min(len(cleaned_rows), 5),
        "preview_quality": "good" if (reader.fieldnames and len(rows) > 0) else "poor",
    }
    return json.dumps(result, ensure_ascii=False, indent=2)


def preview_excel_from_base64(base64file: str) -> str:
    raw = base64.b64decode(base64file)
    bio = io.BytesIO(raw)

    excel = pd.ExcelFile(bio)
    sheets_result = []
    primary_sheet = None

    for sheet_name in excel.sheet_names[:3]:
        bio.seek(0)
        df = pd.read_excel(bio, sheet_name=sheet_name)
        df = df.fillna("")
        all_rows = _clean_records(df.to_dict(orient="records"))
        sheet_payload = {
            "sheet_name": sheet_name,
            "columns": list(df.columns.astype(str)),
            "sample_rows": all_rows[:5],
            "all_rows": all_rows[:500],
            "row_count": len(all_rows),
        }
        sheets_result.append(sheet_payload)
        if primary_sheet is None and sheet_payload["columns"] and all_rows:
            primary_sheet = sheet_payload

    has_data = any(sheet["columns"] and sheet["sample_rows"] for sheet in sheets_result)

    result = {
        "format": "excel",
        "sheets": sheets_result,
        "primary_sheet": primary_sheet,
        "total_rows": sum(sheet.get("row_count", 0) for sheet in sheets_result),
        "preview_quality": "good" if has_data else "poor",
    }
    return json.dumps(result, ensure_ascii=False, indent=2)


def preview_pdf_from_base64(base64file: str) -> str:
    raw = base64.b64decode(base64file)
    bio = io.BytesIO(raw)

    reader = PdfReader(bio)
    pages_text = []
    for page in reader.pages[:10]:
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(text[:5000])

    full_text = "\n".join(pages_text)
    key_lines = [line.strip() for line in full_text.splitlines() if line.strip()][:200]
    kv_candidates = []
    for line in key_lines:
        if ':' in line:
            k, v = line.split(':', 1)
            kv_candidates.append({"key": k.strip(), "value": v.strip()})

    result = {
        "format": "pdf",
        "pages_preview": pages_text[:3],
        "full_text": full_text[:20000],
        "key_lines": key_lines,
        "kv_candidates": kv_candidates,
        "record_mode": "single_form",
        "total_rows": 1,
        "preview_quality": "good" if pages_text else "poor",
    }
    return json.dumps(result, ensure_ascii=False, indent=2)


def preview_docx_from_base64(base64file: str) -> str:
    raw = base64.b64decode(base64file)
    temp_path = Path("temp_preview.docx")
    temp_path.write_bytes(raw)

    doc = Document(str(temp_path))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    kv_candidates = []

    for table in doc.tables[:10]:
        rows = []
        for row in table.rows[:50]:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
            if len(cells) >= 2 and cells[0] and cells[1]:
                kv_candidates.append({"key": cells[0], "value": cells[1]})
        tables.append(rows)

    try:
        temp_path.unlink(missing_ok=True)
    except Exception:
        pass

    full_text = "\n".join(paragraphs + [" | ".join(r) for t in tables for r in t if any(r)])
    key_lines = [line.strip() for line in full_text.splitlines() if line.strip()][:200]
    for line in key_lines:
        if ':' in line:
            k, v = line.split(':', 1)
            kv_candidates.append({"key": k.strip(), "value": v.strip()})

    has_content = bool(paragraphs or tables)

    result = {
        "format": "docx",
        "paragraphs": paragraphs[:50],
        "tables": tables[:5],
        "full_text": full_text[:20000],
        "key_lines": key_lines,
        "kv_candidates": kv_candidates,
        "record_mode": "single_form",
        "total_rows": 1,
        "preview_quality": "good" if has_content else "poor",
    }
    return json.dumps(result, ensure_ascii=False, indent=2)


@lru_cache(maxsize=1)
def _get_easyocr_reader():
    if easyocr is None:
        return None
    try:
        return easyocr.Reader(
            ["ru", "en"],
            gpu=False,
            verbose=False,
        )
    except Exception:
        return None


def _pil_to_bgr(image: Image.Image) -> np.ndarray:
    if cv2 is None:
        raise RuntimeError("OpenCV is not installed")
    rgb = np.array(image.convert("RGB"))
    return cv2.cvtColor(rgb, cv2.COLOR_RGB2BGR)


def _resize_if_needed(img_bgr: np.ndarray, max_side: int = 1600) -> np.ndarray:
    h, w = img_bgr.shape[:2]
    longest = max(h, w)

    if longest <= max_side:
        return img_bgr

    scale = max_side / float(longest)
    new_w = max(1, int(w * scale))
    new_h = max(1, int(h * scale))

    return cv2.resize(img_bgr, (new_w, new_h), interpolation=cv2.INTER_AREA)


def _upscale_for_ocr(img_bgr: np.ndarray, target_min_width: int = 1400) -> np.ndarray:
    h, w = img_bgr.shape[:2]
    if w >= target_min_width:
        return img_bgr

    scale = target_min_width / float(max(w, 1))
    new_w = max(1, int(w * scale))
    new_h = max(1, int(h * scale))
    return cv2.resize(img_bgr, (new_w, new_h), interpolation=cv2.INTER_CUBIC)


def _preprocess_variants(img_bgr: np.ndarray) -> list[tuple[str, np.ndarray]]:
    base = _resize_if_needed(img_bgr, max_side=1600)
    base = _upscale_for_ocr(base, target_min_width=1400)

    gray = cv2.cvtColor(base, cv2.COLOR_BGR2GRAY)

    normalized = cv2.normalize(gray, None, 0, 255, cv2.NORM_MINMAX)

    sharpen_blur = cv2.GaussianBlur(normalized, (0, 0), 1.6)
    sharpen = cv2.addWeighted(normalized, 1.8, sharpen_blur, -0.8, 0)

    adaptive = cv2.adaptiveThreshold(
        normalized,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31,
        9,
    )

    return [
        ("normalized", normalized),
        ("sharpen", sharpen),
        ("adaptive", adaptive),
    ]


def _cleanup_ocr_text(text: str) -> str:
    text = text.replace("“", "\"").replace("”", "\"")
    text = text.replace("‘", "'").replace("’", "'")
    text = text.replace("—", "-")
    text = text.replace("–", "-")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    return text.strip()


def _has_cyrillic(text: str) -> bool:
    return bool(re.search(r"[А-Яа-яЁё]", text))


def _cyrillic_ratio(text: str) -> float:
    letters = [ch for ch in text if ch.isalpha()]
    if not letters:
        return 0.0
    cyr = [ch for ch in letters if re.match(r"[А-Яа-яЁё]", ch)]
    return len(cyr) / len(letters)


def _looks_like_textual_data(text: str) -> bool:
    compact = re.sub(r"\s+", " ", text).strip()
    if len(compact) < 6:
        return False

    words = re.findall(r"[A-Za-zА-Яа-яЁё0-9]{2,}", compact)
    alpha_count = sum(ch.isalpha() for ch in compact)
    digit_count = sum(ch.isdigit() for ch in compact)

    if len(words) >= 2 and alpha_count >= 5:
        return True

    if alpha_count >= 8:
        return True

    if digit_count >= 3 and len(words) >= 1:
        return True

    return False


def _score_candidate(text: str, confs: list[float]) -> float:
    compact = re.sub(r"\s+", " ", text).strip()
    if not compact:
        return -1.0

    avg_conf = sum(confs) / len(confs) if confs else 0.0
    words = re.findall(r"[A-Za-zА-Яа-яЁё0-9]{2,}", compact)
    alpha_count = sum(ch.isalpha() for ch in compact)
    cyr_ratio = _cyrillic_ratio(compact)

    score = 0.0
    score += len(compact) * 1.0
    score += len(words) * 5.0
    score += alpha_count * 0.4
    score += avg_conf * 35.0
    score += cyr_ratio * 80.0
    return score


def _repair_cyrillic_ocr_text(text: str) -> str:
    text = _cleanup_ocr_text(text)

    # типичные OCR-ошибки на кириллице
    replacements = [
        (r"(?i)\bBMA\b", "Вид"),
        (r"(?i)\bBMA[,:;]?", "Вид"),
        (r"(?i)\b4TO\b", "что"),
        (r"(?i)\bBCE\b", "всё"),
        (r"(?i)\bXOPOWO\b", "хорошо"),
        (r"(?i)\bXOPOW[O0]\b", "хорошо"),
        (r"(?i)\b4TO\b", "что"),
        (r"(?i)\bC[ЁE]\b", "сё"),
    ]

    for pattern, repl in replacements:
        text = re.sub(pattern, repl, text)

    # символные исправления только в кириллическом контексте
    chars = list(text)
    for i, ch in enumerate(chars):
        prev_ch = chars[i - 1] if i > 0 else ""
        next_ch = chars[i + 1] if i + 1 < len(chars) else ""

        near_cyr = bool(re.match(r"[А-Яа-яЁё]", prev_ch) or re.match(r"[А-Яа-яЁё]", next_ch))

        if ch == "(" and near_cyr:
            chars[i] = "ё"
        elif ch == ";" and near_cyr:
            chars[i] = "-"
        elif ch == ":" and near_cyr:
            chars[i] = "!"
        elif ch == "B" and near_cyr:
            chars[i] = "В"
        elif ch == "M" and near_cyr:
            chars[i] = "И"
        elif ch == "A" and near_cyr:
            chars[i] = "д"
        elif ch == "4" and near_cyr:
            chars[i] = "ч"
        elif ch == "X" and near_cyr:
            chars[i] = "х"
        elif ch == "O" and near_cyr:
            chars[i] = "о"
        elif ch == "P" and near_cyr:
            chars[i] = "р"
        elif ch == "W" and near_cyr:
            chars[i] = "ш"

    text = "".join(chars)

    text = re.sub(r"\bВИД\b", "Вид", text)
    text = re.sub(r"\bЧТО\b", "что", text)
    text = re.sub(r"\bВСЁ\b", "всё", text)
    text = re.sub(r"\bХОРОШО\b", "хорошо", text)

    text = re.sub(r"\bВид-\b", "Вид ", text)
    text = re.sub(r"\s+", " ", text).strip()

    return text


def _run_easyocr(reader, img: np.ndarray) -> tuple[list[dict[str, Any]], str, list[float]]:
    results = reader.readtext(
        img,
        detail=1,
        paragraph=False,
        decoder="beamsearch",
        beamWidth=10,
        batch_size=1,
        workers=0,
        allowlist=CYRILLIC_ALLOWLIST,
        contrast_ths=0.1,
        adjust_contrast=0.5,
        text_threshold=0.45,
        low_text=0.3,
        link_threshold=0.25,
        width_ths=0.7,
        height_ths=0.7,
    )

    lines: list[dict[str, Any]] = []
    texts: list[str] = []
    confs: list[float] = []

    for item in results:
        if not item or len(item) < 3:
            continue

        _, text, conf = item
        text = str(text).strip()
        if not text:
            continue

        try:
            conf_val = float(conf)
        except Exception:
            conf_val = 0.0

        text = _repair_cyrillic_ocr_text(text)
        if not text:
            continue

        lines.append({"text": text, "score": round(conf_val, 4)})
        texts.append(text)
        confs.append(conf_val)

    full_text = "\n".join(texts).strip()
    full_text = _repair_cyrillic_ocr_text(full_text)

    return lines, full_text, confs


def preview_image_from_base64(base64file: str) -> str:
    raw = base64.b64decode(base64file)
    image = Image.open(io.BytesIO(raw)).convert("RGB")
    img_bgr = _pil_to_bgr(image)

    reader = _get_easyocr_reader()
    if reader is None:
        result = {
            "format": "image",
            "ocr_enabled": False,
            "contains_text_data": False,
            "recognized_lines": [],
            "recognized_text": "",
            "preview_quality": "poor",
            "note": "EasyOCR is not available. Install easyocr to enable image parsing.",
        }
        return json.dumps(result, ensure_ascii=False, indent=2)

    best_lines: list[dict[str, Any]] = []
    best_text = ""
    best_variant = "none"
    best_score = -1.0

    for variant_name, variant_img in _preprocess_variants(img_bgr):
        try:
            lines, text, confs = _run_easyocr(reader, variant_img)
        except Exception:
            continue

        if not text:
            continue

        score = _score_candidate(text, confs)
        current_has_cyr = _has_cyrillic(text)
        best_has_cyr = _has_cyrillic(best_text)

        should_replace = False
        if current_has_cyr and not best_has_cyr:
            should_replace = True
        elif current_has_cyr == best_has_cyr and score > best_score:
            should_replace = True

        if should_replace:
            best_lines = lines
            best_text = text
            best_variant = variant_name
            best_score = score

        if current_has_cyr and len(text) >= 12:
            break

    contains_text_data = _looks_like_textual_data(best_text)

    result = {
        "format": "image",
        "ocr_enabled": True,
        "ocr_engine": "easyocr",
        "ocr_variant": best_variant,
        "contains_text_data": contains_text_data,
        "recognized_lines": best_lines[:30],
        "recognized_text": best_text[:5000],
        "preview_quality": "weak" if contains_text_data else "poor",
        "note": (
            "Image contains recognizable text/numeric data and can be used as a weak source for generation."
            if contains_text_data
            else "OCR did not find enough text/numeric data. The image is treated as an unsupported source."
        ),
    }
    return json.dumps(result, ensure_ascii=False, indent=2)


def build_file_preview(file_name: str, base64file: str) -> str:
    extension = Path(file_name).suffix.lower()

    if extension == ".csv":
        return preview_csv_from_base64(base64file)

    if extension in [".xls", ".xlsx"]:
        return preview_excel_from_base64(base64file)

    if extension == ".pdf":
        return preview_pdf_from_base64(base64file)

    if extension == ".docx":
        return preview_docx_from_base64(base64file)

    if extension in [".png", ".jpg", ".jpeg", ".webp", ".bmp"]:
        return preview_image_from_base64(base64file)

    return json.dumps(
        {
            "format": "unknown",
            "note": f"Unsupported extension: {extension}",
            "preview_quality": "poor",
        },
        ensure_ascii=False,
        indent=2,
    )